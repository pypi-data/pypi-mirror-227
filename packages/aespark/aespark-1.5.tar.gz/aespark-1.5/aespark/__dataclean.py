import re
from docx import Document
from pathlib import Path
from tqdm import tqdm
import pandas as pd
import time
import numpy as np
import warnings
warnings.filterwarnings('ignore')

def dc_dochide(docxurl:str, nameurl:str='', many:bool=False):
    '''
    功能简介：
        word文档脱敏；
    参数解释：
        docxurl 需要脱敏的文档路径
        nameurl 需要脱敏的名字txt文件路径，不传入则仅对数字脱敏
        many 是否要一次性操作多个docx文档，默认false，如果设置为true，docxurl需要传入文件夹路径
    文档解释：
        需要一个记录需要脱敏名称的txt文件，多个名字用空格隔开或者换行的方式写入；
    '''
    if nameurl != '':
        nlist = []
        name = open(nameurl,'r').readlines()
        for i in name:
            i = i.replace('\n','')
            nlist.append(i.split())
        nlist = [item for sublist in nlist for item in sublist]
        nlist = list(set(nlist))

    num_regex = re.compile(r'\d{10,}')

    def tuomin(url):
        doc = Document(url)
        for i in range(len(doc.paragraphs)):
            text = doc.paragraphs[i].text
            if nameurl != '':
                for x in nlist:
                    new_name = x[0] + '某' + x[-1]
                    text = text.replace(x, new_name)
            for num_match in num_regex.finditer(text):
                start = num_match.start()
                end = num_match.end()
                text = text[:start+3] + '8'*(end-start-7) + text[end-4:]
            doc.paragraphs[i].text = text
        for i in range(len(doc.tables)):
            for j in range(len(doc.tables[i].rows)):
                for x in range(len(doc.tables[i].rows[j].cells)):
                    text = doc.tables[i].rows[j].cells[x].text
                    if nameurl != '':
                        for l in nlist:
                            new_name = l[0] + '某' + l[-1]
                            text = text.replace(l, new_name)
                    for num_match in num_regex.finditer(text):
                        start = num_match.start()
                        end = num_match.end()
                        text = text[:start+3] + '8'*(end-start-7) + text[end-4:]
                    doc.tables[i].rows[j].cells[x].text = text

        doc.save(url[:url.rindex('.')]+'-脱敏完成'+url[url.rindex('.'):])

    if many:
        docxs = [str(i) for i in Path(docxurl).rglob('*.docx')]  # 获取目标文件夹下的所有文件路径
        for i in tqdm(docxs, desc='文档脱敏'):
            tuomin(i)
    else:
        tuomin(docxurl)

def dc_write(coun):
    '''格式化金额(不成功则返回原始内容)'''
    try:
        coun = format(float('{:.2f}'.format(float(coun))), ',')
        coun[coun.index('.')+1:]
        coun = coun if len(coun[coun.index('.')+1:]) == 2 else coun+'0'
    finally:
        return coun

def dc_exceladdtt(df: pd.DataFrame):
    '''
    功能简介：
        添加'\\t'便于存为csv文件；
    '''
    df.columns = [str(i)+'\t' if str(i).isdigit()
                  and len(str(i)) > 15 else i for i in df.columns]
    for i in df.columns:
        df[i] = df[i].apply(lambda x: str(
            x)+'\t' if str(x).isdigit() and len(str(x)) > 15 else x)

    return df

def dc_csvdeltt(df: pd.DataFrame):
    '''
    功能简介：
        删除'\\t'便于后续操作；
    '''
    df.columns = [str(i).replace('\t', '') for i in df.columns]
    for i in df.columns:
        df[i] = df[i].apply(lambda x: str(x).replace('\t', ''))

    return df

def dc_invischardel(chars: str | pd.DataFrame, df: bool = False):
    '''
    功能简介：
        清除不可见字符；
    参数解释：
        chars 可传字符可传表，默认传的字符；
        df 如果要传dataframe，该项参数需要填写为 True；
    '''
    if df:
        for i in chars.columns:
            chars[i] = chars[i].apply(lambda x: re.sub(
                u'[\u2000-\u200f\u2028-\u202f\u205f-\u206e]', '', x) if type(x) == str else x)

        return chars
    else:
        if type(chars) == str:
            chars = re.sub(
                u'[\u2000-\u200f\u2028-\u202f\u205f-\u206e]', '', chars)

        return chars
    
def dc_amount(df:pd.DataFrame, clo:str):
    '''
    功能简介：
        交易金额清洗；
    阐述解释：
        df 需要清洗的表；
        clo 需要清洗的列名；
    '''
    try:
        df[clo] = df[clo].astype('float')
    except:
        df.reset_index(drop=True, inplace=True)
        count = 0
        for i in df.index:
            try:
                float(df[clo][i])
            except:
                df[clo][i] = np.nan
                count += 1

        df[clo] = df[clo].astype('float')
        print(f"金额清洗：共清洗错误金额{count}条，占比：{'{:.2%}'.format(count/len(df))}")
    finally:
        return df

def dc_inandout(str: str):
    '''
    功能简介：
        统一借贷标志；
    所需参数：
        需要清洗的字符，建议配合pandas.apply使用；
    当前可清洗内容：
        出 = ['借', '出', '支出', '付', 'D']；
        进 = ['贷', '进', '收入', '收', 'C']；
    如果发现了新的借贷标志可以进行添加；
    '''
    jie = ['借', '出', '支出', '付', 'D']
    dai = ['贷', '进', '收入', '收', 'C']
    if str in jie:
        return '出'
    if str in dai:
        return '进'
    return '其他'

def __try(timestr: str, format: str):
    '''
    功能简介：
        格式化时间格式；
    所需参数：
        timestr 需要格式化的字符串；
        format 字符串的格式（%Y年 %m月 %d日 %H时 %M分 %S秒）；
    return：
        清洗成功的时间格式（示例 2023.07.25 16:11:52）；
        若清洗失败则会返回False；
    '''
    timestr = str(timestr)
    try:
        timeStruct = time.strptime(timestr, format)
        times = time.strftime("%Y.%m.%d %H:%M:%S", timeStruct)
        return times
    except:
        return False

def dc_time(timestr: str):
    '''
    功能简介：
        兼容格式，批量格式化时间格式；
    所需参数：
        timestr 需要格式化的字符串（建议配合pandas.apply使用）；
    return：
        清洗成功的时间格式（示例 2023.07.25 16:11:52）；
        若清洗失败则会返回 nan；
    '''
    timestr = str(timestr)
    if timestr.isdigit():
        if len(timestr) == 14:
            times = __try(timestr, '%Y%m%d%H%M%S')
        elif len(timestr) == 12:
            times = __try(timestr, '%Y%m%d%H%M')
        elif len(timestr) == 8:
            times = __try(timestr, '%Y%m%d')
        else:
            times = __try(timestr, '%Y%m%d%H%M%S')
            if times is False:
                times = __try(timestr, '%Y%m%d%H%M')
            if times is False:
                times = __try(timestr, '%Y%m%d')

    else:
        if '-' in timestr:
            s = '-'
        elif '/' in timestr:
            s = '/'
        elif '.' in timestr:
            s = '.'
        else:
            s = ''
        times = __try(timestr, f'%Y{s}%m{s}%d %H:%M:%S')

        if times is False:
            times = __try(timestr, f'%Y{s}%m{s}%d %H:%M')

        if times is False:
            times = __try(timestr, f'%Y{s}%m{s}%d')

        if times is False and len(timestr) == 26:  # 2016-01-21-21.17.03.704713
            times = __try(timestr[:-7], '%Y-%m-%d-%H.%M.%S')

    return times if times else np.nan