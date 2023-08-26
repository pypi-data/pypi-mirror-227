from .utils import transliterate
from .mail_sender import Mail_sender

from docx import Document
from docx.shared import Inches

from IPython.display import clear_output

import pandas as pd
import torch
import numpy as np

import matplotlib.pyplot as plt

import os
import re
import difflib # �������� ���������� ����� ��������



class Obraz(): # Mail_sender
  def __init__(self, obraz=None, lang="cs", book=None, adress_book=None, model_chekpoint='sergiyvl/model_65000_20ep'):
    """
    ����� ��� ������ � �������.
    obraz - str, ��� �����. ���������� ��� �������� �������. ����� ������ ������ (�� ����� ������).
    book - [" ", " "]����� ������ �����, � ����� �������� ��������.
    adress_book - ������������ ����������� �� ������ (�������������� ��� ������) ��������� �����.

    ������:
    analize - �������� �������. � ��� ���������� ������ book.

    """
    # super().__init__("smtp.mail.ru", "le_i_van@mail.ru", "9dx-U7V-KJr-b2H")
    self.obraz = self._set_obraz(obraz)
    if lang == "greek":
      book = "all_alt_greek_15082023.csv"
      self.lang = "greek"
    else:
      book = "mto_148000_with_adress_newtriodi_23082023.csv"
      self.lang = "cs"
    self._wholebook = None
    self.book = None
    self.email = None
    self.model_chekpoint = model_chekpoint
    self.model = None
    self.tokenizer = None
    self.key_tokens = None # �������� ������. ������ ��������
    self.key_symbols = None # �������� ��������������. ������ �����
    self._book_helpfull = None
    self.texts_with_key_key = None # ������ ����������, ��� ����������� �������� ��������������
    self.texts_without_key_key = None # ������ ����������, ��� ��� �������� ��������������
    self._fast_print_book = None # ������ �� ���������� �������� ��������

    if adress_book != None:
      self.add_book(adress_book=adress_book)
    elif book != None:
      self.add_book(book=book)

#########################################################################################################################################################
################################################### �������� ��� ������, ������� ����� ������������ (������ ���� ����, ��� ����������� ����������� ������� ���������)

  def analize(self):
    """
    �������� �������. �������� �� �� ������������ ���� book, ������� �������� � ��������� ���. ���� �� ��� �� ��������, �� ��� �������� ��� �������.
    """
    if self.book == None:
      raise ValueError("�������� ����� ��� �������. ��� ����� ����� ��������������� ������� .add_book()")


    if self.model == None or self.tokenizer == None:
      self.set_model_and_tokenizer()

    print("��� �����: \""+self.obraz+"\"")
    if self.key_tokens == None:
      self.set_key_tokens()
    if self.key_symbols == None:
      self.set_key_symbols()

    print("�����������... ������ ����� �������� �� 15 �����.")
    self._result = self._analize()

    self._make_book_helpfull(self._result)

    self.texts_with_key_key = []
    self.texts_without_key_key = []
    for text in self._book_helpfull:
        key_key_exist = True
        for key_key in self.key_symbols:
            if key_key.upper() not in text[0].upper():
                key_key_exist = False
        if key_key_exist:
            self.texts_with_key_key.append(text)
        else:
            self.texts_without_key_key.append(text)

    i = 0
    print("������ ������ � ��������� ���������. ��� ������ ����� ����� �������, ���� ������� ����� .results(): ")
    for text in self.texts_with_key_key:
        i += 1
        if i < 40:
          print(i, text)

    # self.send_to_email()

    self.working_with_bookhelpfull()

################################################### ������������ ������� �����

  def fast_or(self):
    """ � self.book ���������� ����� �������������� � ������, ��� ���� ��������������, �� ������ �����������. ��������� - ���������
        ��� ������� ��� ������������� ���������� ��������� � ��� ����������� ������ � ������, �������� ��� ����� �����
    """
    keys = self._input_for_fast()
    i = 0
    self._fast_book = []
    for text in self.book:
      y = False
      for key in keys:
        if key.upper() in text[1].upper():
          y = True
        # for word in text.split():
        #   if key.upper() in word.upper():
        #     print(word)
      if y:
        i += 1
        self._fast_book.append(text)
    self._fast_print_book = self._glue_together_similar_texts(self._fast_book)
    self.book = self._fast_book
    print("� ���������� ", i, " �������, ��� ���� ���� �� ���� �� �������������� ", str(keys))
    self._fast_print(keys, 'or')

  def fast_and(self):
    """ � self.book ���������� ����� �������������� � ������, ��� ���� ��������������, �� ������ �����������. ��������� - ���������
        ��� ������� ��� ������������� ���������� ��������� � ��� ����������� ������ � ������, �������� ��� ����� �����
    """
    keys = self._input_for_fast()
    i = 0
    self._fast_book = []
    for text in self.book:
      y = True
      for key in keys:
        if key.upper() not in text[1].upper():
          y = False
        # for word in text.split():
        #   if key.upper() in word.upper():
        #     print(word)
      if y:
        i += 1
        self._fast_book.append(text)
    self._fast_print_book = self._glue_together_similar_texts(self._fast_book)
    self.book = self._fast_book
    print("� ���������� ", i, " �������, ��� � ������ ���� ��� �������������� ", str(keys))
    self._fast_print(keys, 'and')

  def fast_or_re(self):
    """ � self.book ���������� ����� �������������� � ������, ��� ���� ��������������, �� ������ �����������. ��������� - ���������
        ����������� ���� ������� � ������������� ���������� ��������� � � ����������� ������ � ������, �������� ��� ����� �����
    """
    keys = self._input_for_fast()
    i = 0
    self._fast_book = []
    pattern = r''
    for key in keys:
      key = key.upper()
      if key[0] == "_":
        key = r'[^�-��-�]'+key[1:]
      if key[-1] == "_":
        key = key[:-1]+r'[^�-��-�]'
      key += '|'
      pattern += key
    pattern = pattern[:-1]

    for text in self.book:
      if re.search(pattern, " "+text[1].upper()+" "):
        i += 1
        self._fast_book.append(text)
    self._fast_print_book = self._glue_together_similar_texts(self._fast_book)
    self.book = self._fast_book
    print("� ���������� ", i, " �������, ��� ���� ���� �� ���� �� �������������� ", str(keys))
    self._fast_print(keys, 'or')


  def fast_and_re(self):
    """ � self.book ���������� ����� �������������� � ������, ��� ���� ��������������, �� ������ �����������. ��������� - ���������
        ����������� ���� ������� � ������������� ���������� ��������� � � ����������� ������ � ������, �������� ��� ����� �����
    """
    keys = self._input_for_fast()
    i = 0
    self._fast_book = []
    patterns = []
    for key in keys:
      key = key.upper()
      if key[0] == "_":
        key = r'[^�-��-�]'+key[1:]
      if key[-1] == "_":
        key = key[:-1]+r'[^�-��-�]'
      patterns.append(key)

    for text in self.book:
      y = True
      for pattern in patterns:
        if not re.search(pattern, " "+text[1].upper()+" "):
          y = False
      if y:
        i += 1
        self._fast_book.append(text)
    self._fast_print_book = self._glue_together_similar_texts(self._fast_book)
    self.book = self._fast_book
    print("� ���������� ", i, " �������, ��� � ������ ���� ��� �������������� ", str(keys))
    self._fast_print(keys, 'and')


#########################################################################################################################################################
################################################### �������� �������, ������� ���� �� ������� ������� �� �����

  def results(self):
    """
    ������� ��� ������ ����������� ������.
    """
    i = 0
    for text in self.texts_with_key_key[:30]:
      i += 1
      print(i, text)
    i = 0
    for text in self.texts_without_key_key[:30]:
      i += 1
      print(i, text)


  def add_book(self, book=None, adress_book=None):
    """�������� ����� ��� ������. ���� �� ������ ����� - ���������� ����� �� ���������.
    """

    if book != None:
      p = os.path.join(os.path.dirname(__file__), '')
      p = os.path.join(p, book)
      csv_book = pd.read_csv(p)

      self._wholebook = [sent[1:4] for sent in csv_book.values]
      for i in range(len(self._wholebook)):
        if type(self._wholebook[i][0]) != type(148.97):
          self._wholebook[i][0] = re.sub(r'<.*?>', '', self._wholebook[i][0])
          self._wholebook[i][0] = re.sub(r'[\n\t]', '', self._wholebook[i][0])
          self._wholebook[i][0] = self._wholebook[i][0].strip()
          self._wholebook[i][1] = re.sub(r'<.*?>', '', self._wholebook[i][1])
          self._wholebook[i][1] = re.sub(r'[\n\t]', '', self._wholebook[i][1])
          self._wholebook[i][1] = self._wholebook[i][1].strip()
          self._wholebook[i][2] = re.sub(r'<.*?>', '', self._wholebook[i][2])
          self._wholebook[i][2] = re.sub(r'[\n\t]', '', self._wholebook[i][2])
          self._wholebook[i][2] = self._wholebook[i][2].strip()
        else:
          self._wholebook[i][0] = "������ �����"
      self.book = self._wholebook

    elif adress_book != None:
      csv_book = pd.read_csv(adress_book)

      self._wholebook = [sent[1:4] for sent in csv_book.values]
      for i in range(len(self._wholebook)):
        if type(self._wholebook[i][0]) != type(148.97):
          self._wholebook[i][0] = re.sub(r'<.*?>', '', self._wholebook[i][0])
          self._wholebook[i][0] = re.sub(r'[\n\t]', '', self._wholebook[i][0])
          self._wholebook[i][0] = self._wholebook[i][0].strip()
          self._wholebook[i][1] = re.sub(r'<.*?>', '', self._wholebook[i][1])
          self._wholebook[i][1] = re.sub(r'[\n\t]', '', self._wholebook[i][1])
          self._wholebook[i][1] = self._wholebook[i][1].strip()
          self._wholebook[i][2] = re.sub(r'<.*?>', '', self._wholebook[i][2])
          self._wholebook[i][2] = re.sub(r'[\n\t]', '', self._wholebook[i][2])
          self._wholebook[i][2] = self._wholebook[i][2].strip()
        else:
          self._wholebook[i][0] = "������ �����"
      self.book = self._wholebook
      # print(self.book[:100])

  def send_to_email(self, email=None):
    """
    ����������� ����������� ��������� ���������� �� �����. �� ���� ����������� �����.
    """
    h_m_with_keykey = int(input("������� ������ ����������� �� ������ ��������? ������� �����: "))
    if self.email != None:
      email = self.email
    elif email == None:
      # inp = input("�� ������ �������� ������ ���������� �� ����� � ���� word �����? ������� �� ��� ���: ")
      # if "���".upper() in inp.upper():
      #   return
      # if input("������ ������� ��������� ������������� ��� ������� � ��������� ��������? ������� 1 ��� 0 ") == "1":
      #   self.show_diagramm(self.texts_with_key_key)
      print("������ �� �����, �������� ��� ������� gmail. yandex ��� outlook ���� ������. mail - �����. �� �������� ��������� ����!")
      email = input("������� ����� ����� ����������� �����: ")
    if self.texts_with_key_key == None:
      print("����� ���������� ���������, �� ����� ��������. ��������� ������ �����. .analize()")
      return
    inp = input("�� ������ �������� ���������� � ������� docx ��� ������ ������ ������? ������� 'docx' ��� ����� �������: ")
    if "docx" in inp:
      self._send_to_email_docx(email, h_m_with_keykey)
    else:
      self._send_to_email(email, h_m_with_keykey)


  def working_with_bookhelpfull(self):
    """
    ������� � ����� ������ � bookhelpfull.
    � � ���� ����� ����� ������ �������������� � ��� ����������� ������������ ������ �� ���� �����, ��� ���� ���� �� ���� �����.
    � ������� ����� ���������� ����� ������ ������ �������.
    """
    print("\n\n����� ������ � bookhelpfull �����������. \n")
    self.book = self._book_helpfull


  def working_with_wholebook(self):
    """
    ������� � ����� ������ �� ���� ������. � � ���� ����� ����� �������������� �� ���� ����������� �����.
    """
    print("\n\n����� ������ �� ���� ������. \n")
    self.book = self._wholebook

  def _glue_together_similar_texts(self, array_of_texts): # �� ���� ����� ��� ������ ������� ����: (�����, ����� ������)
    ans = input("����� ����������. ���������� "+ str(len(array_of_texts)) + " �������. �� ����� ����� ������������ ������� ����� ����� ���������� � ���������� � �������� ������ �� ������.\n" +
                "����������, ��� 10 ������� ������������� �� ���� ������, 100 �������� �� 40, � 500 �� ��������� �����. \n" +
                "������� \"��\", ���� ������ ��� �������, ����� ������� ����� �������: ")
    if ans.strip().upper() == "��": 
      table_of_simil = []
      for text_1 in array_of_texts:
        texts_simil = []
        for text_2 in array_of_texts:
          texts_simil.append(difflib.SequenceMatcher(None, text_1[1], text_2[1]).ratio()) #  real_quick_ratio() quick_ratio() ratio()
        table_of_simil.append(texts_simil)

      array_in_work = [{"text_el": array_of_texts[i], "base_id": i} for i in range(len(array_of_texts))]
      
      if len(table_of_simil) > 0:
        for i in range(len(table_of_simil[0])):
          for j in range(i, len(table_of_simil)):
            if i != j: # ����� �� ���������!!
              if table_of_simil[i][j] >= 0.85:
                array_in_work[j]["base_id"] = array_in_work[i]["base_id"]

      simil_itog = [[] for _ in array_in_work]
      for el in array_in_work:
        if len(simil_itog[el["base_id"]]) == 0: # ���� ����� ������� ��� ������� base_id � ������� simil_itog ����� ����, ������, � ���� ������ ������ �� ���������, ������, ���� ������ ���������
          simil_itog[el["base_id"]].append(el["text_el"][0])
          simil_itog[el["base_id"]].append(el["text_el"][1])
          simil_itog[el["base_id"]].append([])
        simil_itog[el["base_id"]][2].append(el["text_el"][2])

      simil_itog_itog = []
      for el in simil_itog:
        if len(el) > 0:
          a = [len(el[2]), el[0], el[1], el[2]]
          simil_itog_itog.append(a)
      return sorted(simil_itog_itog, reverse=True)
    else: 
      arr = []
      i = 0
      for el in array_of_texts: 
        i += 1
        arr.append([str(i), el[0], el[1], el[2]])
      return arr

  def _altgreek_diakr_to_lite(self, text):
    """ ��������� ������ ��������� ������� � ����������� � ��������� ����� ��� ����������
    """
    dict_?ltgreek_diakr = {
      "???????????????????????????????????????????????": "?",
      "?????????????????": "?",
      "???????????????????????????????????????????": "?",
      "?????????????????????????????": "?",
      "?????????????????": "?",
      "???????????????????????????????????????????": "?",
      "?????????????????????????": "?",
      "?'???": "?",
      "?`???????????????": "",
      "?,.?": " "
    }

    text = list(text)
    for i in range(len(text)):
      for key in dict_?ltgreek_diakr.keys():
        if text[i] in key:
          text[i] = dict_?ltgreek_diakr[key]

    return "".join(text)

  def _input_for_fast(self):
    """������� � ��������� ������ �� ��������� ��������
    """
    text = input("����� Fast. ������� ����� ������ �� ��������������, ������� ������ ���� � ������������� ������: ")
    if self.lang == "greek":
      text = self._altgreek_diakr_to_lite(text)
    return [key for key in text.split()]



#########################################################################################################################################################
################################################### set ������. ��� ��������� �������, ������� ������ ������� �������. ��������, ��� ����� � ����� ��� ������

  def set_key_tokens(self):
    """
    ������ �������� ������.
    """
    if self.tokenizer != None:
      obraz_tokenize = self.tokenizer.tokenize(self.obraz)

      i = 0
      for token in obraz_tokenize:
          i += 1
          print("(", i,", "+token+")", end="  ", sep="")
          if i%10 == 0:
              print()
      self.key_tokens = [int(i) for i in input("\n������� ����� ������ ������ �������� �������: ").split()]

  def set_key_symbols(self):
    """������ �������� ��������������."""
    self.key_symbols = [i for i in input("������� ��������������, ��� ������� � ������ ����� ��� ������.\n���� �������� �������� ���������, ����� ������� ����� ������:").split()]

  def set_model_and_tokenizer(self, model_chekpoint=None, model=None, tokenizer=None):
    print("��������� ������ ��� �������...")
    self.device = torch.device("cuda:0" if torch.cuda.is_available() else "cpu")

    if model_chekpoint != None:
      self.model_chekpoint = model_chekpoint

    if model == None:
      from transformers import AutoTokenizer
      self.tokenizer = AutoTokenizer.from_pretrained(self.model_chekpoint)

      from transformers import AutoModelForMaskedLM
      self.model = AutoModelForMaskedLM.from_pretrained(self.model_chekpoint)
    else:
      self.model = model
      self.tokenizer = tokenizer

    self.model.to(self.device)
    clear_output(True)
    clear_output(True)
    clear_output(True)

  def _set_obraz(self, obraz):
    obraz = input("������� �����: ")
    while True:
      try:
        if type(obraz) != type("str"):
          raise TypeError("� �������� ������ ���������� �������� ������ (str).")
      except TypeError:
        print("� �������� ������ ���������� �������� ������ (str).")
        obraz = input("������� �����: ")
      else:
        break
    if obraz[-1] != ".":
      obraz += "."
    clear_output()
    print("��� �����: \""+obraz+"\"")

    return obraz

#########################################################################################################################################################
################################################### ������� ������ ������ � �������������� ��������� ����� (���� ��� ������ ������� �������)

  def _analize(self):
    """
    ������� ��������. ������������ �� ������� ������. � ��������� ���� ��� ������� ����� �������, ����� � ����������.
    �������� ������ ���������� � �������. �������� ������, ���������� �����, ������� ���� �� �����. (otn = batch_max[i][ind-1]/batch_max[i][ind]).
    ���� ����������� ������ �������� ��������������.
    ��������� ������ ������� �� ��������� �������� �������. (���������� - ����� ������� > 200)
    """

    obraz_tokenize = self.tokenizer.tokenize(self.obraz)

    relevant = [(1 if i+1 not in self.key_tokens else 1.5) for i in range(len(obraz_tokenize))]


    # ������������ ����� � ����������� �����
    batch_size=64
    text_tokenized = self.tokenizer(self.obraz, return_tensors='pt').to(self.device)
    text_throw_model = self.model(
                            **text_tokenized,
                            output_hidden_states=True
                        ).hidden_states[:][-1]


    words = text_throw_model[0][1:-1].reshape(1, -1, 768)

    normalization_token = []
    for word in obraz_tokenize:
        token = self.tokenizer(word+'.', return_tensors='pt').to(self.device)
        token_1 = self.model(
            **token,
            output_hidden_states=True
        ).hidden_states[:][-1]
        token_2 = torch.transpose(token_1, 1, 2)
        koef = 40/(token_1 @ token_2)[0, 1, 1].detach().cpu().numpy()
        normalization_token.append(koef)

    normalization_word = []
    # normilization = 40/(words@torch.transpose(words, 1, 2)) ��� �� ���� ����� ����������
    for word in words[0]:
        word_transpose = torch.transpose(word.reshape(1, -1), 0, 1)
        koef = 40/(word.reshape(1, -1) @ word_transpose)[0, 0].detach().cpu().numpy()
        normalization_word.append(koef)

    if False:
      for i in range(len(normalization_word)):
          print(obraz_tokenize[i], normalization_token[i], normalization_word[i])

    # ���� ���� �������� �������, ������� �������� �����, ����� � ������� ������������� ������������ �� ���������� �������������
    pairs_second = []
    for i in range(len(relevant)):
        if relevant[i] > 1:
            if obraz_tokenize[i][:2] == "##":
                if relevant[i-1] > 1:
                    pairs_second.append(i)

    # ��������� ����� ������ � ��������� � result �����
    result = np.zeros((len(self.book), 2, len(relevant)))

    self.local_book = [text[0] for text in self.book]

    for start_index in range(0, len(self.local_book), batch_size):
        batch = self.local_book[start_index:start_index+batch_size]

        batch = self.tokenizer(batch, return_tensors='pt',truncation=True, padding=True, max_length=45).to(self.device)

        batch = self.model(
        **batch,
        output_hidden_states=True
        ).hidden_states[:][-1]

    #  ������ ���� ������������� � �������� ������ � ������ ���������� �����������
        batch = torch.transpose(batch, 1, 2)

        batch = (words @ batch).detach().cpu().numpy()
        batch = batch[:, :, 1:-1]

        batch_max = np.max(batch, axis=2)
        for i in np.arange(batch_size):
            if start_index+i<len(self.local_book):
                result[start_index+i][0][0] = start_index+i
                for ind in pairs_second:
                  otn = batch_max[i][ind-1]/batch_max[i][ind]
                  normalization_word[ind-1] *= (otn if otn < 1 else 1/otn)

                result[start_index+i][1] = (batch_max[i]*normalization_word)**relevant # ������ ����� ������������
                # ��� ������ �������� ������� ��� �����. ������� ����� ����� �����������!!! ������ ����� �������, �� ����� ��������� � bool_helpfull
                koef = sum((result[start_index+i][1] > 200)*1)
                result[start_index+i][1] *= (koef+1)/2

    return result


#########################################################################################################################################################
################################################### �������� ���������� ��� ����������� �������� �� �����

  def make_result_dokument_not_docx(self, file_name, title=None, first_paragraf=None, h_m_with_keykey=70, h_m_without_keykey=20): #  adress = "/content/drive/MyDrive/diplom/results/results/",

    result_doc = """"""

    # document = Document()

    if title != None:
      result_doc += title.upper() + "\n" + "\n"

    # document.add_heading('�������� ������� ������', level=1)
    if first_paragraf != None:
      result_doc += first_paragraf + "\n"

    obraz_tokenized = self.tokenizer.tokenize(self.obraz)

    result_doc += "\n" + "\n" + "\n" + '������ ������ � ��������� ����������������'.upper() + "\n" + "\n"
    result_doc += "�������� ��������������: " + str(self.key_symbols)
    i = 0
    for text in self.texts_with_key_key[:h_m_with_keykey]:
      i += 1
      result_doc += str(i)+" "+text + "\n"

    result_doc += "\n" + "\n" + "\n" + '������ ������ ��� �������� ��������������'.upper()  + "\n" + "\n"
    i = 0
    for text in self.texts_without_key_key[:h_m_without_keykey]:
      i += 1
      result_doc += str(i)+" "+text + "\n"


    # document.save(adress + file_name)
    return result_doc


  def _make_result_dokument_docx(self, file_name, title=None, first_paragraf=None, h_m_with_keykey=70, h_m_without_keykey=20, adress = '/content/'): #  adress = '/content/',

    document = Document()

    if title != None:
      document.add_heading(title, 0)

    document.add_heading('�������� ������� ������', level=1)
    if first_paragraf != None:
      p = document.add_paragraph(first_paragraf)

    obraz_tokenized = self.tokenizer.tokenize(self.obraz)

    document.add_heading('������ ������ � ��������� ����������������', level=1)
    document.add_paragraph("�������� ��������������: " + str(self.key_symbols))

    i = 0
    num = input("�������� ���������? ������� �� ��� ���").strip().upper()
    for text in self.texts_with_key_key[:h_m_with_keykey]:
      if num == "���".upper():
        p = document.add_paragraph(text[0]+"  ")
      else:
        i += 1
        p = document.add_paragraph(str(i)+" "+text[0]+"  ")
      p.add_run("("+text[2]+")").italic = True

    document.add_heading('������ ������ ��� �������� ��������������', level=1)
    i = 0
    for text in self.texts_without_key_key[:h_m_without_keykey]:
      i += 1
      p = document.add_paragraph(str(i)+" "+text[0]+"  ")
      p.add_run("("+text[2]+")").italic = True



    document.save(adress + file_name)


  def _make_fast_book_dokument_docx(self, file_name, keys, title=None, mode=None, adress = '/content/'): #  adress = '/content/',

    document = Document()

    if title != None:
      document.add_heading(title, 0)

    document.add_heading('������ ������ �������:', level=1)
    if mode == 'or':
      p = document.add_paragraph("������, ��� ���� ���� �� ���� �� �������������� " + str(keys)[1:-1])
    elif mode == 'and':
      p = document.add_paragraph("������, ��� � ������ ���� ��� �������������� " + str(keys)[1:-1])

    i = 0
    num = input("�������� ���������? ������� �� ��� ���").strip().upper()
    for text in self._fast_print_book:
      if num == "���".upper():
        p = document.add_paragraph(text[0]+"  ")
      else:
        i += 1
        p = document.add_paragraph(str(i)+". "+str(text[0])+" - "+text[1]+"   ")
      p.add_run("("+", ".join(text[2])+")").italic = True


    document.save(adress + file_name)




#########################################################################################################################################################
 ################################################### ����� ������� � �������� �� �����

  def _fast_print(self, keys, mode):
    choice = int(input("������ ������� ��������� ������? ������� '0', ���� �� ������. \n" +
          "������� '1', ���� ������ ������� ��������� � �������. \n" +
          "������� '2', ���� ������ �������� ��������� �� �����. \n" +
          "������� ����� (� ������� Enter): " ))
    if choice == 0:
      return
    if choice == 1:
      for text in self._fast_print_book:
        print(text[0], ", "+text[1]+",    ("+", ".join(text[3])+")", sep='')
    if choice == 2:
      if self.email != None:
        email = self.email
      else:
        email = input("������� ��� email: ")
      self._send_to_email_docx_fast_book(keys_=keys, mode=mode, email=email)


  # def _show(self, my_list, h_m=30):
  #   print("������ "+str(h_m)+" ���������, �������������")

  #   plt.plot(my_list[:h_m])
  #   plt.grid()
  #   plt.show()


  def _make_book_helpfull(self, result):
    self._book_helpfull = []

    result_list = list(result)
    result_list.sort(key=lambda x: sum(x[1]), reverse=True)
    i = 0
    for el in result_list[:4000]:
        i += 1
        self._book_helpfull.append(self.book[int(el[0, 0])])

  def _send_to_email(self, email, h_m_with_keykey):
    obraz_tokens, obraz_key_tokens = ["", ""]

    obraz_tokenize = self.tokenizer.tokenize(self.obraz)

    for i in range(1, len(obraz_tokenize)):
      obraz_tokens += "'"+obraz_tokenize[i-1]+"' "
      if i in self.key_tokens:
        obraz_key_tokens += "'"+obraz_tokenize[i-1]+"' "

    # h_m = len()
    name_of_file = transliterate(obraz_key_tokens) + ".docx"
    par = """������ ����������� ��� ������: '""" + self.obraz + """'
    ����������� ��������� �� ������: """ + obraz_tokens + """
    �������� ������: """ + obraz_key_tokens + str(self.key_tokens) +"""
    �������� ��������������: """ + str(self.key_symbols) + """
    model: '"""+ self.model_chekpoint +"""'\n
    ����: ���� ������ �������� � �� ��������. �������� ������� � ������� 1.5, � �� �������� � ������� 1. ��� ���������� ����������� ��� ������
    ����� ������������� ������������ ����� ���� ������ ������. �� ���� ������� ������ ������� ������, � ������������� (@) � ����� �����. ��������� ������������ - 40/���������� ��� ������������ �����.
    ����������� �������� ����, ������� ��������� �� ������ (##). �������������� ����������� - ��������� �������� �� ���� � ��������. �������� ������ ��� �������� �������.
    ������� ����� �� �������� ��������.  ��������� ������ ������� �� ��������� �������� �������. (����������� - ����� ������� > 150)

    """
    document = self.make_result_dokument_not_docx(name_of_file, title=str(self.obraz), first_paragraf=par, h_m_with_keykey=h_m_with_keykey)

    self.send_mail("���������� ��� "+self.obraz[:17]+"", document, email)

  def _send_to_email_docx(self, email, h_m_with_keykey):
    obraz_tokens, obraz_key_tokens = ["", ""]

    obraz_tokenize = self.tokenizer.tokenize(self.obraz)

    for i in range(1, len(obraz_tokenize)):
      obraz_tokens += "'"+obraz_tokenize[i-1]+"' "
      if i in self.key_tokens:
        obraz_key_tokens += "'"+obraz_tokenize[i-1]+"' "

    # h_m = len()
    name_of_file = transliterate(obraz_key_tokens) + ".docx"
    par = """������ ����������� ��� ������: '""" + self.obraz + """'
    ����������� ��������� �� ������: """ + obraz_tokens + """
    �������� ������: """ + obraz_key_tokens + str(self.key_tokens) +"""
    �������� ��������������: """ + str(self.key_symbols) + """
    model: '"""+ self.model_chekpoint +"""'\n
    ����: ���� ������ �������� � �� ��������. �������� ������� � ������� 1.5, � �� �������� � ������� 1. ��� ���������� ����������� ��� ������
    ����� ������������� ������������ ����� ���� ������ ������. �� ���� ������� ������ ������� ������, � ������������� (@) � ����� �����. ��������� ������������ - 40/���������� ��� ������������ �����.
    ����������� �������� ����, ������� ��������� �� ������ (##). �������������� ����������� - ��������� �������� �� ���� � ��������. �������� ������ ��� �������� �������.
    ������� ����� �� �������� ��������.  ��������� ������ ������� �� ��������� �������� �������. (����������� - ����� ������� > 150)

    """
    self._make_result_dokument_docx(name_of_file, title=str(self.obraz), first_paragraf=par, h_m_with_keykey=h_m_with_keykey)

    body_text = "���������� ��� "+self.obraz+"" + "\n���������� ������������ � � � ... � ��������� ��� ����."

    self.send_mail("���������� ��� "+self.obraz[:17]+"", body_text, email, file_to_attach='/content/'+name_of_file)
    self.send_mail(email+self.obraz[:17]+"", "��������� �� �����: " + email + "\n" + body_text, 'vl.sergiiy@gmail.com', file_to_attach='/content/'+name_of_file)



  def _send_to_email_docx_fast_book(self, keys_, mode, email):

    keys = ""
    for k in keys_:
      keys += " " + k

    name_of_file = mode+" " + transliterate(keys) + ".docx"

    self._make_fast_book_dokument_docx(name_of_file, keys, title=mode+" "+keys, mode=mode)

    body_text = "��� ����������. "+mode+keys

    self.send_mail("���������� ��� "+keys+"", body_text, email, file_to_attach='/content/'+name_of_file)
    self.send_mail(email+keys+"", "��������� �� �����: " + email + "\n" + body_text, 'vl.sergiiy@gmail.com', file_to_attach='/content/'+name_of_file)


#########################################################################################################################################################
################################################### ����� �� ���������� (�� ��������� �������)

  def show_diagramm(self, array_for_diagramm=None): # �� ���� �� ��������
    """���������
    """
    if array_for_diagramm == None:
      i = 0
      what_array = {}
      print("��������� ������ ������� �� ������ �������? ")
      if self.texts_with_key_key != None:
        print(i, "��������������� �����, ��� ����� ������ �� ���������� ������ ������� analize, � ������� ���� �������� ��������������")
        what_array[i] = self.texts_with_key_key
        i += 1
      if self.texts_without_key_key != None:
        print(i, "��������������� �����, ��� ����� ������ �� ���������� ������ ������� analize, � ������� ��� �������� ��������������")
        what_array[i] = self.texts_without_key_key
        i += 1
      if self._book_helpfull != None:
        print(i, "��������������� �����, ��� ����� ��� ���������� ������ ������� analize, � ������� ���� ���� �� ���� �������� �����")
        what_array[i] = self._book_helpfull
        i += 1
      if self._wholebook != None:
        print(i, "��������� �����")
        what_array[i] = self._wholebook
        i += 1
      ind = int(input("������� ����� �������: "))
      array_for_diagramm = what_array[ind]


    my_list = []
    array = self._result[0][1] > 1
    array = self._result[0][1][array]
    i = 0
    for el in self._result:
      if self._wholebook[int(el[0][0])] == array_for_diagramm[i]:
        my_list.append(int(sum(el[1])/(5*array.shape[0])))
        i += 1
    my_list.sort(reverse=True)


    h_m = []
    for el in [10500, 1000, 500, 70, 30, len(my_list)]:
      if el <= len(my_list):
        h_m.append(el)
    for el in sorted(h_m):
      self._show(my_list, el)



